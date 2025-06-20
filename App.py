
import streamlit as st
import yfinance as yf
import pandas as pd
import matplotlib.pyplot as plt
from statsmodels.tsa.arima.model import ARIMA
from sklearn.metrics import mean_squared_error, mean_absolute_error
import numpy as np
from datetime import datetime
import io
from docx import Document
from fpdf import FPDF

# App Title
st.title("📈 Stock Market Forecasting using ARIMA")
st.write("This app forecasts stock prices using the ARIMA model and allows downloads.")

# Sidebar Inputs
st.sidebar.header("🔍 Stock Selection")
stock_symbol = st.sidebar.text_input("Enter Stock Symbol (e.g., AAPL, MSFT, TCS.NS):", "AAPL")
start_date = st.sidebar.date_input("Start Date", datetime(2015, 1, 1))
end_date = st.sidebar.date_input("End Date", datetime.today())
forecast_days = st.sidebar.slider("Forecast days:", min_value=7, max_value=90, value=30)

# ARIMA Parameters
st.sidebar.subheader("⚙️ ARIMA Parameters")
p = st.sidebar.number_input("p (AR term)", min_value=0, max_value=10, value=5)
d = st.sidebar.number_input("d (Differencing)", min_value=0, max_value=2, value=1)
q = st.sidebar.number_input("q (MA term)", min_value=0, max_value=10, value=0)

# Load Data
@st.cache_data
def load_data(symbol, start, end):
    data = yf.download(symbol, start=start, end=end)
    if isinstance(data.columns, pd.MultiIndex):
        data.columns = data.columns.get_level_values(0)
    data = data[['Close']]
    data.dropna(inplace=True)
    return data

data_load_state = st.text("Loading data...")
df = load_data(stock_symbol, start_date, end_date)
data_load_state.text("Loading data...done!")

# Show Historical Data
st.subheader("📉 Historical Close Price")
st.line_chart(df)

# Fit ARIMA Model
st.subheader("🔮 ARIMA Forecasting")
try:
    model = ARIMA(df['Close'], order=(p, d, q))
    model_fit = model.fit()
    forecast = model_fit.forecast(steps=forecast_days)
    forecast_index = pd.date_range(start=df.index[-1], periods=forecast_days + 1, freq='B')[1:]
    forecast_df = pd.DataFrame(forecast, index=forecast_index, columns=['Forecast'])

    # Plot Forecast
    fig, ax = plt.subplots(figsize=(10, 4))
    df['Close'].plot(ax=ax, label='Historical')
    forecast_df['Forecast'].plot(ax=ax, label='Forecast', color='orange')
    plt.title(f"{stock_symbol} Forecast for Next {forecast_days} Days")
    plt.xlabel("Date")
    plt.ylabel("Price")
    plt.legend()
    st.pyplot(fig)

    st.write("Forecast Data")
    st.dataframe(forecast_df)

    # Accuracy Evaluation
    actual = df['Close'][-forecast_days:]
    common_index = actual.index.intersection(forecast_df.index)
    actual_aligned = actual.loc[common_index]
    forecast_aligned = forecast_df.loc[common_index]['Forecast']

    if not actual_aligned.empty:
        rmse = np.sqrt(mean_squared_error(actual_aligned, forecast_aligned))
        mae = mean_absolute_error(actual_aligned, forecast_aligned)
        st.subheader("📏 Model Accuracy")
        st.write(f"**RMSE:** {rmse:.2f}")
        st.write(f"**MAE:** {mae:.2f}")
    else:
        st.info("Not enough data overlap for accuracy evaluation.")

    # Download: CSV
    csv = forecast_df.to_csv().encode('utf-8')
    st.download_button("⬇️ Download CSV", csv, "forecast.csv", "text/csv")

    # Download: PDF
    def create_pdf(dataframe):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Stock Forecast Report", ln=True, align='C')
        for i, (date, row) in enumerate(dataframe.iterrows()):
            pdf.cell(200, 10, txt=f"{date.strftime('%Y-%m-%d')} : {row['Forecast']:.2f}", ln=True)
        buffer = io.BytesIO()
        pdf.output(buffer, 'S')
        buffer.seek(0)
        return buffer

    pdf_file = create_pdf(forecast_df)
    st.download_button("⬇️ Download PDF", pdf_file, "forecast.pdf", "application/pdf")

    # Download: DOCX
    def create_docx(dataframe):
        doc = Document()
        doc.add_heading("Stock Forecast Report", 0)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Forecast'
        for date, row in dataframe.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(date.date())
            row_cells[1].text = f"{row['Forecast']:.2f}"
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    docx_file = create_docx(forecast_df)
    st.download_button("⬇️ Download DOCX", docx_file, "forecast.docx",
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

except Exception as e:
    st.error(f"Error: {e}")
